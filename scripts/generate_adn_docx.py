"""
Script para generar documento DOCX profesional con información de pruebas de ADN
Laboratorios Pérez - Potosí, Bolivia
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Rutas
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
BASE_DIR = os.path.dirname(SCRIPT_DIR)  # Go up one level from 'scripts' folder
LOGO_PATH = os.path.join(BASE_DIR, 'app', 'static', 'img', 'logo.jpg')
OUTPUT_PATH = os.path.join(BASE_DIR, 'app', 'static', 'uploads', 'docs', 'pruebas_adn_perez.docx')

def set_cell_shading(cell, color):
    """Aplica color de fondo a una celda"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def create_adn_document():
    doc = Document()
    
    # ========== CONFIGURAR MÁRGENES ==========
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # ========== ENCABEZADO CON LOGO ==========
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = True
    
    # Logo (columna izquierda)
    logo_cell = header_table.rows[0].cells[0]
    logo_para = logo_cell.paragraphs[0]
    logo_run = logo_para.add_run()
    if os.path.exists(LOGO_PATH):
        logo_run.add_picture(LOGO_PATH, width=Inches(1.5))
    
    # Título (columna derecha)
    title_cell = header_table.rows[0].cells[1]
    title_para = title_cell.paragraphs[0]
    title_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    title_run = title_para.add_run("LABORATORIOS PÉREZ")
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(26, 188, 156)  # Verde teal
    
    title_para.add_run("\n")
    subtitle_run = title_para.add_run("Laboratorio Clínico • Potosí, Bolivia")
    subtitle_run.font.size = Pt(10)
    subtitle_run.font.color.rgb = RGBColor(100, 100, 100)
    
    # Línea separadora
    doc.add_paragraph()
    separator = doc.add_paragraph()
    sep_run = separator.add_run("─" * 70)
    sep_run.font.color.rgb = RGBColor(26, 188, 156)
    separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ========== TÍTULO PRINCIPAL ==========
    main_title = doc.add_paragraph()
    main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = main_title.add_run("PRUEBA DE ADN")
    title_run.bold = True
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(44, 62, 80)
    
    doc.add_paragraph()
    
    # ========== INTRODUCCIÓN ==========
    intro = doc.add_paragraph()
    intro_run = intro.add_run("Gracias por escribirnos a Laboratorios Pérez, con relación a la prueba de paternidad por ADN le enviamos la siguiente información.")
    intro_run.font.size = Pt(11)
    
    doc.add_paragraph()
    
    # ========== OPCIONES DE PRUEBA ==========
    options_title = doc.add_paragraph()
    opt_run = options_title.add_run("Para acceder a la prueba de paternidad existen dos opciones:")
    opt_run.bold = True
    opt_run.font.size = Pt(12)
    opt_run.font.color.rgb = RGBColor(26, 188, 156)
    
    opt1 = doc.add_paragraph(style='List Bullet')
    opt1.add_run("La primera con la presencia del Padre, de la Madre y del Niñ@.")
    
    opt2 = doc.add_paragraph(style='List Bullet')
    opt2.add_run("La segunda es solo con la presencia del Padre y del Niñ@.")
    
    doc.add_paragraph()
    
    # ========== SOBRE LAS MUESTRAS (TABLA) ==========
    muestras_title = doc.add_paragraph()
    m_run = muestras_title.add_run("SOBRE LAS MUESTRAS")
    m_run.bold = True
    m_run.font.size = Pt(14)
    m_run.font.color.rgb = RGBColor(243, 156, 18)  # Naranja
    
    # Tabla de muestras
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    
    # Encabezados
    headers = ["Tipo de Muestra", "Descripción", "Costo", "Tiempo"]
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, '1ABC9C')  # Verde teal
    
    # Datos de muestras
    muestras_data = [
        ["💉 Sangre (micro-muestra)", "6 gotitas de muestra sanguínea a partir de un pequeño pinchazo en la pulpa del dedo. Recolección en oficinas de Laboratorios Pérez.", "2500 Bs", "7 días hábiles"],
        ["🦷 Hisopado Bucal", "Tomado del Padre y del Niñ@. Recolección en oficinas de Laboratorios Pérez.", "2500 Bs", "7 días hábiles"],
        ["💇 Cabello + Sangre", "10 cabellitos del Niñ@ arrancados (con raíz, no cortados) + muestra sanguínea del Padre.", "3500 Bs", "7 días hábiles"],
        ["💅 Uñitas + Sangre", "5 pedazos de uñitas del Niñ@ (mano o pie) + muestra sanguínea del Padre.", "3500 Bs", "7 días hábiles"],
    ]
    
    for row_idx, row_data in enumerate(muestras_data, start=1):
        row = table.rows[row_idx]
        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = cell_text
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
            # Color alterno de filas
            if row_idx % 2 == 0:
                set_cell_shading(cell, 'F8F9FA')
    
    doc.add_paragraph()
    
    # ========== HORARIOS ==========
    horario_title = doc.add_paragraph()
    h_run = horario_title.add_run("HORARIOS DE ATENCIÓN")
    h_run.bold = True
    h_run.font.size = Pt(12)
    h_run.font.color.rgb = RGBColor(26, 188, 156)
    
    h1 = doc.add_paragraph(style='List Bullet')
    h1.add_run("Lunes a Viernes: ").bold = True
    h1.add_run("7:30 a.m. a 6:00 p.m.")
    
    h2 = doc.add_paragraph(style='List Bullet')
    h2.add_run("Sábados: ").bold = True
    h2.add_run("7:30 a.m. a 12:00 p.m.")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # ========== PIE DE PÁGINA ==========
    footer_line = doc.add_paragraph()
    footer_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fl_run = footer_line.add_run("─" * 70)
    fl_run.font.color.rgb = RGBColor(26, 188, 156)
    
    thanks = doc.add_paragraph()
    thanks.alignment = WD_ALIGN_PARAGRAPH.CENTER
    thanks_run = thanks.add_run("Quedamos a su disposición para cualquier duda adicional.\n¡Gracias por confiar en nuestro laboratorio!")
    thanks_run.italic = True
    thanks_run.font.size = Pt(10)
    thanks_run.font.color.rgb = RGBColor(100, 100, 100)
    
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_run = contact.add_run("📞 67619188  •  📧 laboratorios.perez@gmail.com")
    contact_run.font.size = Pt(9)
    contact_run.font.color.rgb = RGBColor(26, 188, 156)
    
    # ========== GUARDAR DOCUMENTO ==========
    doc.save(OUTPUT_PATH)
    print(f"✅ Documento creado exitosamente: {OUTPUT_PATH}")
    return OUTPUT_PATH

if __name__ == "__main__":
    create_adn_document()
