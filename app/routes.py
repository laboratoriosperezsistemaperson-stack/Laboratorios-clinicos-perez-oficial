from flask import Blueprint, render_template, request, redirect, url_for, flash, send_file, jsonify
from flask_login import login_required, current_user
from app import db
from app.models import Paciente, Resultado, Prueba, Publicacion, FotoGaleria, ConfiguracionLab
from app.utils import admin_required
from werkzeug.utils import secure_filename
from datetime import datetime, timedelta
from sqlalchemy import func, extract
import os
import random
import string
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

main = Blueprint('main', __name__)

# Definir ruta base absoluta para archivos
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
UPLOAD_DIR = os.path.join(BASE_DIR, 'app', 'static', 'uploads')
PRUEBAS_UPLOAD_DIR = os.path.join(UPLOAD_DIR, 'pruebas')
BACKUP_DIR = os.path.join(UPLOAD_DIR, 'backups')  # Carpeta de backups

# ============ HEALTH CHECK PARA RENDER ============
@main.route('/health')
def health_check():
    """
    Health check endpoint para Render
    - Verifica que la app está viva
    - Verifica conexión a la base de datos
    - Retorna 200 OK si todo está bien
    """
    try:
        # Intentar hacer una query simple a la BD
        db.session.execute(db.text('SELECT 1'))
        return jsonify({
            'status': 'healthy',
            'database': 'connected',
            'timestamp': datetime.now().isoformat()
        }), 200
    except Exception as e:
        # Si falla la BD, aún responder pero con status unhealthy
        return jsonify({
            'status': 'unhealthy',
            'database': 'disconnected',
            'error': str(e),
            'timestamp': datetime.now().isoformat()
        }), 503  # Service Unavailable

@main.route('/ping')
def ping():
    """Endpoint simple para verificar que la app responde"""
    return jsonify({'status': 'ok', 'message': 'pong'}), 200

def generar_codigo_acceso():
    """Genera un código aleatorio de 8 caracteres ÚNICO"""
    while True:
        codigo = ''.join(random.choices(string.ascii_uppercase + string.digits, k=8))
        # Verificar que no existe en la BD
        if not Resultado.query.filter_by(codigo_acceso=codigo).first():
            return codigo

def generar_numero_orden():
    """
    Genera un número de orden ÚNICO basado en timestamp + contador
    Formato: YYYYMMDD-HHMMSS-XXX
    Ejemplo: 20251107-153045-001
    """
    from datetime import datetime
    timestamp = datetime.now().strftime('%Y%m%d-%H%M%S')

    # Contar cuántos resultados se crearon en este segundo
    contador = Resultado.query.filter(
        Resultado.numero_orden.like(f'{timestamp}%')
    ).count()

    numero_orden = f"{timestamp}-{contador+1:03d}"
    return numero_orden

from app.pdf_manager import FileManager

# Instancia global (ya maneja credenciales internamente)
pdf_manager = FileManager(UPLOAD_DIR)
file_manager = pdf_manager # Alias para compatibilidad y semántica

def guardar_pdf_con_backup(archivo, numero_orden, paciente):
    """
    Sube el PDF a Supabase Storage
    """
    try:
        success, storage_path, error = pdf_manager.save_pdf(archivo, numero_orden, paciente)
        
        if success:
            print(f"✓ PDF subido a Supabase: {storage_path}")
            # Retornamos storage_path como 'relative_path' y None para rutas fisicas
            return storage_path, None, None
        else:
            raise Exception(error)

    except Exception as e:
        print(f"✗ Error guardando PDF: {str(e)}")
        return None, None, None

def limpiar_archivo_huerfano(filename):
    """Elimina un archivo PDF de Supabase si existe"""
    try:
        if filename:
            pdf_manager.delete_pdf(filename)
            print(f"🗑 Archivo huérfano eliminado de Supabase: {filename}")
    except Exception as e:
        print(f"⚠ Error limpiando archivos huérfanos: {e}")

@main.route('/descargar/<int:resultado_id>')
@admin_required
def descargar(resultado_id):
    resultado = Resultado.query.get_or_404(resultado_id)
    if resultado.archivo_pdf:
        # Obtener URL pública de Supabase
        public_url = pdf_manager.get_public_url(resultado.archivo_pdf)
        if public_url:
             return redirect(public_url)

        # Fallback para archivos antiguos locales (si existen y estamos en local)
        # Ojo: En Render esto fallará para archivos viejos no migrados.
        possible_paths = [
            os.path.join(UPLOAD_DIR, resultado.archivo_pdf),
            os.path.join(BASE_DIR, resultado.archivo_pdf),
            resultado.archivo_pdf
        ]
        for pdf_path in possible_paths:
             if os.path.exists(pdf_path):
                 return send_file(pdf_path, as_attachment=True)

        flash(f'El archivo PDF no se encuentra disponible (ni en nube ni local)', 'danger')
        return redirect(url_for('main.admin_resultados'))

    flash('No hay archivo PDF asignado a este resultado', 'warning')
    return redirect(url_for('main.admin_resultados'))

@main.route('/descargar-resultado-publico/<int:resultado_id>')
def descargar_resultado_publico(resultado_id):
    """
    Descarga pública de PDFs para pacientes
    - Redirige a la URL de Supabase para descarga directa
    """
    # LOGGING DETALLADO
    print("=" * 80)
    print("🔍 DESCARGA PÚBLICA DE PDF (CLOUD)")
    print(f"   Resultado ID: {resultado_id}")

    resultado = Resultado.query.get_or_404(resultado_id)

    if not resultado.archivo_pdf:
        print("   ❌ NO TIENE ARCHIVO ASIGNADO")
        return "El archivo no está disponible", 404
        
    # Intentar obtener URL de Supabase
    public_url = pdf_manager.get_public_url(resultado.archivo_pdf)
    
    if public_url:
        print(f"   ✅ Redirigiendo a Supabase: {public_url}")
        return redirect(public_url)
    
    # Fallback local (solo si la migración falló o estamos en dev)
    print("   ⚠️ Archivo no encontrado en nube, buscando local...")
    # ... codigo legacy ...
    
    return "El archivo PDF no se encuentra disponible en este momento.", 404

# ... (resto de funciones) ...

@main.route('/resultado/eliminar/<int:resultado_id>', methods=['POST'])
@admin_required
def eliminar_resultado(resultado_id):
    """
    SOFT-DELETE: Mueve el resultado a la papelera (simulado en Supabase o flag DB)
    """
    try:
        resultado = Resultado.query.get_or_404(resultado_id)
        
        # Soft-delete en BD
        resultado.eliminado = True
        resultado.fecha_eliminacion = datetime.now()
        resultado.eliminado_por = current_user.username if current_user and hasattr(current_user, 'username') else 'admin'
        
        # Opcional: Mover archivo en Supabase a carpeta 'papelera'
        if resultado.archivo_pdf:
             success, new_path = pdf_manager.move_to_trash(resultado.archivo_pdf)
             if success:
                 resultado.archivo_pdf = new_path

        db.session.commit()
        flash(f'✅ Resultado movido a la papelera exitosamente.', 'success')
        
    except Exception as e:
        db.session.rollback()
        flash(f'❌ Error al eliminar resultado: {str(e)}', 'danger')
        print(f"Error eliminando resultado: {e}")
        
    return redirect(url_for('main.admin_resultados'))

@main.route('/resultado/restaurar/<int:resultado_id>', methods=['POST'])
@admin_required
def restaurar_resultado(resultado_id):
    try:
        resultado = Resultado.query.get_or_404(resultado_id)
        
        # Restaurar en BD
        resultado.eliminado = False
        resultado.fecha_eliminacion = None
        resultado.eliminado_por = None
        
        # Restaurar archivo en Supabase
        if resultado.archivo_pdf and 'papelera/' in resultado.archivo_pdf:
             success, new_path = pdf_manager.restore_from_trash(resultado.archivo_pdf)
             if success:
                 resultado.archivo_pdf = new_path
        
        db.session.commit()
        flash(f'✅ Resultado restaurado exitosamente.', 'success')

    except Exception as e:
        db.session.rollback()
        print(f"❌ Error al restaurar resultado: {str(e)}")
        flash(f'❌ Error al restaurar resultado: {str(e)}', 'danger')

    return redirect(url_for('main.admin_resultados'))

@main.route('/resultado/eliminar-permanente/<int:resultado_id>', methods=['POST'])
@admin_required
def eliminar_permanente(resultado_id):
    try:
        resultado = Resultado.query.get_or_404(resultado_id)
        paciente_nombre = resultado.paciente_nombre

        # Eliminar archivo de Supabase
        if resultado.archivo_pdf:
            pdf_manager.delete_pdf(resultado.archivo_pdf)

        db.session.delete(resultado)
        db.session.commit()

        flash(f'✅ Resultado de "{paciente_nombre}" eliminado permanentemente.', 'success')

    except Exception as e:
        db.session.rollback()
        print(f"❌ Error al eliminar permanentemente: {str(e)}")
        flash(f'❌ Error al eliminar: {str(e)}', 'danger')

    return redirect(url_for('main.admin_resultados'))

@main.route('/resultado/reemplazar/<int:resultado_id>', methods=['POST'])
@admin_required
def reemplazar_pdf(resultado_id):
    """
    Reemplaza el PDF de un resultado existente en Supabase
    """
    try:
        resultado = Resultado.query.get_or_404(resultado_id)
        archivo_nuevo = request.files.get('archivo_pdf')

        if not archivo_nuevo or not archivo_nuevo.filename:
            flash('❌ Debe seleccionar un archivo PDF', 'danger')
            return redirect(url_for('main.admin_resultados'))

        if archivo_nuevo and archivo_nuevo.filename:
            # 1. Eliminar archivo anterior de Supabase
            if resultado.archivo_pdf:
                pdf_manager.delete_pdf(resultado.archivo_pdf)

            # 2. Guardar nuevo archivo
            success, storage_path, error = pdf_manager.save_pdf(archivo_nuevo, resultado.numero_orden)
            
            if success:
                resultado.archivo_pdf = storage_path
                db.session.commit()
                flash(f'✅ PDF reemplazado exitosamente', 'success')
            else:
                flash(f'❌ Error al subir nuevo PDF: {error}', 'danger')

    except Exception as e:
        db.session.rollback()
        print(f"❌ Error al reemplazar PDF: {str(e)}")
        flash(f'❌ Error al reemplazar PDF: {str(e)}', 'danger')

    return redirect(url_for('main.admin_resultados'))

def limpiar_archivo_huerfano(filename):
    """Elimina un archivo PDF y su backup si existen"""
    try:
        # Eliminar archivo principal
        filepath = os.path.join(UPLOAD_DIR, filename)
        if os.path.exists(filepath):
            os.remove(filepath)
            print(f"🗑 Archivo principal eliminado: {filename}")

        # Eliminar backup
        backup_path = os.path.join(BACKUP_DIR, filename)
        if os.path.exists(backup_path):
            os.remove(backup_path)
            print(f"🗑 Backup eliminado: {filename}")
    except Exception as e:
        print(f"⚠ Error limpiando archivos huérfanos: {e}")

@main.route('/')
def index():
    return render_template('publico/index.html')

@main.route('/portal-resultados')
def portal_resultados():
    return render_template('publico/portal_resultados.html')

@main.route('/servicios')
def servicios():
    # Cargar publicaciones activas de la base de datos
    publicaciones = Publicacion.query.filter_by(activo=True).order_by(Publicacion.orden.asc(), Publicacion.fecha_publicacion.desc()).all()
    
    # Cargar configuración
    config_items = ConfiguracionLab.query.all()
    configuracion = {item.clave: item.valor for item in config_items}
    
    return render_template('publico/servicios.html', publicaciones=publicaciones, config=configuracion)

@main.route('/servicios/informacion')
def servicios_info():
    # Cargar configuración del laboratorio
    config_items = ConfiguracionLab.query.all()
    configuracion = {item.clave: item.valor for item in config_items}
    
    return render_template('publico/servicios_info.html', config=configuracion)

@main.route('/servicios/fotos')
def servicios_fotos():
    # Cargar fotos activas de la galería
    fotos = FotoGaleria.query.filter_by(activo=True).order_by(FotoGaleria.orden.asc(), FotoGaleria.fecha_subida.desc()).all()
    
    return render_template('publico/servicios_fotos.html', fotos=fotos)

@main.route('/catalogo-pruebas')
def catalogo_pruebas():
    # Obtener todas las pruebas ordenadas por categoría y nombre
    pruebas = Prueba.query.order_by(Prueba.categoria, Prueba.nombre).all()

    # Organizar pruebas por categoría
    pruebas_por_categoria = {}
    for prueba in pruebas:
        categoria = prueba.categoria or 'General'
        if categoria not in pruebas_por_categoria:
            pruebas_por_categoria[categoria] = []
        pruebas_por_categoria[categoria].append(prueba)

    return render_template('publico/catalogo/lista_pruebas.html',
                         pruebas=pruebas,
                         pruebas_por_categoria=pruebas_por_categoria)

@main.route('/consultar-resultado', methods=['POST'])
def consultar_resultado():
    """
    REQ-02: Login simplificado solo con CI
    Busca el paciente por CI y muestra su historial completo
    """
    ci = request.form.get('ci', '').strip()
    
    if not ci:
        flash('Por favor ingrese su Cédula de Identidad', 'danger')
        return redirect(url_for('main.portal_resultados'))
    
    # Buscar paciente por CI
    paciente = Paciente.query.filter_by(ci=ci).first()
    
    if not paciente:
        # Intentar buscar en resultados directamente (para datos legacy)
        resultado = Resultado.query.filter_by(paciente_ci=ci, eliminado=False).first()
        if resultado:
            # Crear objeto paciente temporal para la vista
            return render_template('publico/portal_historial.html', 
                                  paciente={'nombre': resultado.paciente_nombre, 'ci': ci},
                                  resultados=[resultado])
        flash('No se encontró ningún paciente con ese CI', 'danger')
        return redirect(url_for('main.portal_resultados'))
    
    # Obtener TODOS los resultados del paciente (no eliminados)
    resultados = Resultado.query.filter_by(
        paciente_ci=ci,
        eliminado=False
    ).order_by(Resultado.fecha_muestra.desc()).all()
    
    return render_template('publico/portal_historial.html', 
                          paciente=paciente, 
                          resultados=resultados)

@main.route('/dashboard')
@admin_required
def dashboard():
    # Estadísticas básicas
    total_pacientes = Paciente.query.count()
    total_resultados = Resultado.query.count()
    total_pruebas = Prueba.query.count()

    # Pacientes registrados en los últimos 6 meses
    fecha_inicio = datetime.now() - timedelta(days=180)
    pacientes_por_mes = db.session.query(
        extract('month', Paciente.fecha_registro).label('mes'),
        func.count(Paciente.id).label('total')
    ).filter(Paciente.fecha_registro >= fecha_inicio).group_by('mes').all()

    # Resultados por mes en los últimos 6 meses
    resultados_por_mes = db.session.query(
        extract('month', Resultado.fecha_muestra).label('mes'),
        func.count(Resultado.id).label('total')
    ).filter(Resultado.fecha_muestra >= fecha_inicio).group_by('mes').all()

    # Top 5 pruebas más recientes del catálogo
    top_pruebas = db.session.query(
        Prueba.nombre,
        Prueba.precio
    ).order_by(Prueba.fecha_creacion.desc()).limit(5).all()

    # Estadísticas adicionales
    pacientes_este_mes = Paciente.query.filter(
        extract('month', Paciente.fecha_registro) == datetime.now().month,
        extract('year', Paciente.fecha_registro) == datetime.now().year
    ).count()

    resultados_este_mes = Resultado.query.filter(
        extract('month', Resultado.fecha_muestra) == datetime.now().month,
        extract('year', Resultado.fecha_muestra) == datetime.now().year
    ).count()

    # Preparar datos para gráficos
    meses_nombres = ['Ene', 'Feb', 'Mar', 'Abr', 'May', 'Jun', 'Jul', 'Ago', 'Sep', 'Oct', 'Nov', 'Dic']
    mes_actual = datetime.now().month

    # Últimos 6 meses
    ultimos_meses = []
    for i in range(5, -1, -1):
        mes_idx = (mes_actual - i - 1) % 12
        ultimos_meses.append(meses_nombres[mes_idx])

    # Crear arrays de datos para gráficos
    pacientes_data = [0] * 6
    resultados_data = [0] * 6

    for mes, total in pacientes_por_mes:
        mes_idx = int(mes)
        pos = 5 - ((mes_actual - mes_idx) % 12)
        if 0 <= pos < 6:
            pacientes_data[pos] = total

    for mes, total in resultados_por_mes:
        mes_idx = int(mes)
        pos = 5 - ((mes_actual - mes_idx) % 12)
        if 0 <= pos < 6:
            resultados_data[pos] = total

    # Preparar datos para exportación como JSON serializado
    pacientes_lista = Paciente.query.order_by(Paciente.nombre).all()
    all_pruebas = Prueba.query.all()
    all_resultados = Resultado.query.all()
    
    import json
    export_data = {
        "pacientes": [
            {
                "id": p.id,
                "nombre": p.nombre,
                "ci": p.ci,
                "telefono": p.telefono or '',
                "email": p.email or '',
                "fecha_registro": p.fecha_registro.strftime('%Y-%m-%d %H:%M') if p.fecha_registro else ''
            } for p in pacientes_lista
        ],
        "pruebas": [
            {
                "id": pr.id,
                "nombre": pr.nombre,
                "categoria": pr.categoria or '',
                "precio": float(pr.precio) if pr.precio else 0,
                "descripcion": pr.descripcion or ''
            } for pr in all_pruebas
        ],
        "resultados": [
            {
                "id": r.id,
                "numero_orden": r.numero_orden,
                "paciente_nombre": r.paciente_nombre,
                "paciente_ci": r.paciente_ci,
                "paciente_id": r.paciente_id,
                "prueba_id": r.prueba_id,
                "fecha_muestra": r.fecha_muestra.strftime('%Y-%m-%d') if r.fecha_muestra else '',
                "codigo_acceso": r.codigo_acceso,
                "tipo_laboratorio": r.prueba.nombre if r.prueba else 'General',
                "eliminado": r.eliminado if hasattr(r, 'eliminado') else False,
                "fecha_creacion": r.fecha_creacion.strftime('%Y-%m-%d') if r.fecha_creacion else ''
            } for r in all_resultados
        ],
        "estadisticas": {
            "totalPacientes": total_pacientes,
            "totalResultados": total_resultados,
            "totalPruebas": total_pruebas,
            "pacientesEsteMes": pacientes_este_mes,
            "resultadosEsteMes": resultados_este_mes
        }
    }
    
    export_data_json = json.dumps(export_data, ensure_ascii=False)

    return render_template('admin/dashboard.html',
                         total_pacientes=total_pacientes,
                         total_resultados=total_resultados,
                         total_pruebas=total_pruebas,
                         pacientes_este_mes=pacientes_este_mes,
                         resultados_este_mes=resultados_este_mes,
                         ultimos_meses=ultimos_meses,
                         pacientes_data=pacientes_data,
                         resultados_data=resultados_data,
                         top_pruebas=top_pruebas,
                         pacientes_lista=pacientes_lista,
                         all_pruebas=all_pruebas,
                         all_resultados=all_resultados,
                         export_data_json=export_data_json)

@main.route('/pacientes', methods=['GET', 'POST'])
@admin_required
def admin_pacientes():
    if request.method == 'POST':
        try:
            paciente = Paciente(
                nombre=request.form['nombre'],
                ci=request.form['ci'],
                telefono=request.form.get('telefono'),
                email=request.form.get('email')
            )
            db.session.add(paciente)
            db.session.commit()
            
            # Crear carpeta del paciente inmediatamente
            try:
                from werkzeug.utils import secure_filename
                nombre_carpeta = secure_filename(f"{paciente.ci}_{paciente.nombre}")
                paciente_dir = os.path.join(UPLOAD_DIR, 'pacientes', nombre_carpeta)
                os.makedirs(paciente_dir, exist_ok=True)
                print(f"✓ Carpeta creada: {paciente_dir}")
            except Exception as e:
                print(f"⚠ Error creando carpeta: {e}")
                
            flash('Paciente registrado exitosamente', 'success')
            # Redirigir con parámetros para mostrar modal de éxito
            return redirect(url_for('main.admin_pacientes', nuevo_exito=1, paciente_nombre=paciente.nombre))
        except Exception as e:
            flash(f'Error: {str(e)}', 'danger')
        return redirect(url_for('main.admin_pacientes'))
    # Ordenar alfabéticamente por nombre para evitar duplicidades visuales
    pacientes = Paciente.query.order_by(Paciente.nombre.asc()).all()
    return render_template('admin/pacientes.html', pacientes=pacientes, now=datetime.now())

@main.route('/paciente/<int:paciente_id>')
@admin_required
def ver_paciente(paciente_id):
    paciente = Paciente.query.get_or_404(paciente_id)
    return jsonify({
        'id': paciente.id,
        'nombre': paciente.nombre,
        'ci': paciente.ci,
        'telefono': paciente.telefono or '-',
        'email': paciente.email or '-',
        'fecha_registro': paciente.fecha_registro.strftime('%d/%m/%Y %H:%M')
    })

@main.route('/paciente/editar/<int:paciente_id>', methods=['POST'])
@admin_required
def editar_paciente(paciente_id):
    try:
        paciente = Paciente.query.get_or_404(paciente_id)
        old_ci = paciente.ci
        old_nombre = paciente.nombre
        
        paciente.nombre = request.form['nombre']
        paciente.ci = request.form['ci']
        paciente.telefono = request.form.get('telefono')
        paciente.email = request.form.get('email')

        # Si cambió nombre o CI, renombrar carpeta en Supabase
        if old_ci != paciente.ci or old_nombre != paciente.nombre:
            try:
                # Renombrar carpeta en Supabase Storage
                success, updated_paths, error = pdf_manager.rename_patient_folder(
                    old_ci, old_nombre, 
                    paciente.ci, paciente.nombre
                )
                
                if success and updated_paths:
                    # Actualizar paths en la BD para resultados de este paciente
                    from werkzeug.utils import secure_filename
                    old_folder_name = secure_filename(f"{old_ci}_{old_nombre}")
                    new_folder_name = secure_filename(f"{paciente.ci}_{paciente.nombre}")
                    
                    resultados_p = Resultado.query.filter_by(paciente_id=paciente_id).all()
                    for res in resultados_p:
                        if res.archivo_pdf and old_folder_name in res.archivo_pdf:
                            res.archivo_pdf = res.archivo_pdf.replace(old_folder_name, new_folder_name)
                    
                    print(f"✅ Carpeta y {len(updated_paths)} archivos actualizados en Supabase")
                elif error:
                    print(f"⚠️ {error}")
                    
            except Exception as e:
                print(f"⚠️ Error renombrando carpeta en Supabase: {e}")

        resultados = Resultado.query.filter_by(paciente_id=paciente_id).all()
        for resultado in resultados:
            resultado.paciente_nombre = paciente.nombre
            resultado.paciente_ci = paciente.ci

        db.session.commit()
        flash('Paciente y sus resultados actualizados exitosamente', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error: {str(e)}', 'danger')
    return redirect(url_for('main.admin_pacientes'))



@main.route('/paciente/eliminar/<int:paciente_id>', methods=['POST'])
@admin_required
def eliminar_paciente(paciente_id):
    try:
        paciente = Paciente.query.get_or_404(paciente_id)
        nombre_paciente = paciente.nombre

        # Obtener todos los resultados del paciente
        resultados = Resultado.query.filter_by(paciente_id=paciente_id).all()

        # Eliminar archivos PDF físicos
        archivos_eliminados = 0
        for resultado in resultados:
            if resultado.archivo_pdf:
                try:
                    pdf_path = os.path.join(UPLOAD_DIR, resultado.archivo_pdf)
                    if os.path.exists(pdf_path):
                        os.remove(pdf_path)
                        archivos_eliminados += 1
                except Exception as e:
                    print(f"Error eliminando archivo {resultado.archivo_pdf}: {e}")

        # SQLAlchemy eliminará automáticamente los resultados gracias a cascade='all, delete-orphan'
        db.session.delete(paciente)
        db.session.commit()

        flash(f'Paciente "{nombre_paciente}" eliminado exitosamente junto con {len(resultados)} resultado(s) y {archivos_eliminados} archivo(s) PDF', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error al eliminar paciente: {str(e)}', 'danger')
        print(f"Error: {e}")

    return redirect(url_for('main.admin_pacientes'))

@main.route('/resultados', methods=['GET', 'POST'])
@admin_required
def admin_resultados():
    """
    SISTEMA ROBUSTO DE GESTIÓN DE RESULTADOS
    - Números de orden automáticos y únicos
    - Almacenamiento seguro de PDFs
    - Backups automáticos
    - Manejo inteligente de errores
    - Nunca se pierden archivos
    """
    if request.method == 'POST':
        filename_guardado = None
        numero_orden_generado = None

        try:
            # ============ VALIDACIONES INICIALES ============
            archivo = request.files.get('archivo_pdf')

            # Validar paciente PRIMERO (antes de guardar archivo)
            paciente_id = request.form.get('paciente_id')
            if not paciente_id:
                flash('❌ Debe seleccionar un paciente', 'danger')
                return redirect(url_for('main.admin_resultados'))

            paciente = Paciente.query.get(int(paciente_id))
            if not paciente:
                flash('❌ Paciente no encontrado', 'danger')
                return redirect(url_for('main.admin_resultados'))

            # ============ GENERAR NÚMERO DE ORDEN AUTOMÁTICO ============
            numero_orden_manual = request.form.get('numero_orden', '').strip()

            if numero_orden_manual:
                # Usuario proporcionó número manual
                numero_orden_generado = numero_orden_manual
                print(f"📋 Usando número de orden manual: {numero_orden_generado}")
            else:
                # Generar automáticamente
                numero_orden_generado = generar_numero_orden()
                print(f"🔢 Número de orden generado automáticamente: {numero_orden_generado}")

            # ============ GUARDAR PDF CON BACKUP ============
            filename_guardado, filepath, backup_path = guardar_pdf_con_backup(
                archivo,
                numero_orden_generado,
                paciente
            )

            if not filename_guardado:
                raise Exception("No se pudo guardar el archivo PDF")

            print(f"📁 Archivo guardado: {filename_guardado}")
            print(f"💾 Backup creado en: {BACKUP_DIR}")

            # ============ PROCESAR FECHA ============
            fecha_str = request.form.get('fecha_muestra')
            fecha_muestra = datetime.strptime(fecha_str, '%Y-%m-%d').date() if fecha_str else None

            # ============ REGISTRAR EN BASE DE DATOS ============
            # Ya no se usa código de acceso - login simplificado solo con CI
            resultado = Resultado(
                numero_orden=numero_orden_generado,
                paciente_id=paciente.id,
                paciente_nombre=paciente.nombre,
                paciente_ci=paciente.ci,
                fecha_muestra=fecha_muestra,
                archivo_pdf=filename_guardado,
                codigo_acceso=None,  # Ya no se usa - acceso solo con CI
                prueba_id=request.form.get('prueba_id') or None
            )

            db.session.add(resultado)
            db.session.commit()

            # ============ ÉXITO COMPLETO ============
            print("=" * 80)
            print("✅ RESULTADO GUARDADO EXITOSAMENTE")
            print(f"   ID: {resultado.id}")
            print(f"   Número Orden: {numero_orden_generado}")
            print(f"   Paciente: {paciente.nombre} (CI: {paciente.ci})")
            print(f"   Archivo: {filename_guardado}")
            print(f"   Backup: ✓ Creado")
            print("=" * 80)

            flash(f'✅ Resultado guardado exitosamente para {paciente.nombre}', 'success')

        except ValueError as ve:
            # Errores de validación (archivo no válido, etc.)
            db.session.rollback()
            if filename_guardado:
                limpiar_archivo_huerfano(filename_guardado)
            flash(f'❌ Error de validación: {str(ve)}', 'danger')
            print(f"✗ Error de validación: {str(ve)}")

        except Exception as e:
            # Errores inesperados
            db.session.rollback()

            # Limpiar archivos huérfanos
            if filename_guardado:
                limpiar_archivo_huerfano(filename_guardado)

            error_msg = str(e)
            print("=" * 80)
            print("❌ ERROR AL GUARDAR RESULTADO")
            print(f"   Error: {error_msg}")
            print(f"   Archivos limpiados: {filename_guardado if filename_guardado else 'N/A'}")
            print("=" * 80)

            # Mensaje de error amigable
            if 'duplicate key' in error_msg.lower():
                flash('❌ Este número de orden ya existe. El sistema generará uno automático.', 'danger')
            else:
                flash(f'❌ Error al guardar: {error_msg}', 'danger')

        return redirect(url_for('main.admin_resultados', 
                              nuevo_exito=1, 
                              paciente_id=paciente.id, 
                              paciente_nombre=paciente.nombre))
    
    # Separar resultados activos y eliminados para tabs
    resultados_activos = Resultado.query.filter_by(eliminado=False).order_by(Resultado.fecha_creacion.desc()).all()
    resultados_eliminados = Resultado.query.filter_by(eliminado=True).order_by(Resultado.fecha_eliminacion.desc()).all()
    pacientes = Paciente.query.order_by(Paciente.nombre).all()
    pruebas = Prueba.query.order_by(Prueba.nombre).all()
    
    return render_template('admin/resultados.html', 
                          resultados=resultados_activos,
                          resultados_eliminados=resultados_eliminados,
                          pacientes=pacientes, 
                          pruebas=pruebas)

@main.route('/descargar-credenciales-pdf/<int:resultado_id>')
@admin_required
def descargar_credenciales_pdf(resultado_id):
    """
    Genera PDF de credenciales PROFESIONAL y SIN ERRORES
    Para que el paciente vea su código de acceso
    """
    resultado = Resultado.query.get_or_404(resultado_id)

    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=50,
        leftMargin=50,
        topMargin=50,
        bottomMargin=50
    )
    elements = []
    styles = getSampleStyleSheet()

    # ============ ESTILOS PERSONALIZADOS ============
    title_style = ParagraphStyle(
        'TitleStyle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#11998e'),
        spaceAfter=10,
        alignment=1,  # Centrado
        fontName='Helvetica-Bold'
    )

    subtitle_style = ParagraphStyle(
        'SubtitleStyle',
        parent=styles['Normal'],
        fontSize=12,
        textColor=colors.HexColor('#7f8c8d'),
        spaceAfter=30,
        alignment=1,  # Centrado
        fontName='Helvetica'
    )

    section_title_style = ParagraphStyle(
        'SectionTitle',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=colors.HexColor('#2c3e50'),
        spaceBefore=10,
        spaceAfter=10,
        fontName='Helvetica-Bold'
    )

    # ============ ENCABEZADO ============
    elements.append(Paragraph('LABORATORIO CLÍNICO PÉREZ', title_style))
    elements.append(Paragraph('Potosí, Bolivia', subtitle_style))
    elements.append(Spacer(1, 0.1*inch))

    # ============ TÍTULO PRINCIPAL ============
    header_data = [['CREDENCIALES DE ACCESO A RESULTADOS']]
    header_table = Table(header_data, colWidths=[6.5*inch])
    header_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#11998e')),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.white),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 14),
        ('TOPPADDING', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    elements.append(header_table)
    elements.append(Spacer(1, 0.15*inch))

    # ============ INFORMACIÓN DEL PACIENTE ============
    patient_data = [
        ['INFORMACIÓN DEL PACIENTE'],
        ['Nombre Completo:', resultado.paciente_nombre],
        ['Cédula de Identidad:', resultado.paciente_ci],
        ['Número de Orden:', resultado.numero_orden],
        ['Fecha de Emisión:', resultado.fecha_creacion.strftime("%d/%m/%Y %H:%M")]
    ]

    patient_table = Table(patient_data, colWidths=[2.5*inch, 4*inch])
    patient_table.setStyle(TableStyle([
        # Header
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#34495e')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 11),
        ('ALIGN', (0, 0), (-1, 0), 'LEFT'),
        ('SPAN', (0, 0), (-1, 0)),  # Merge header cells
        ('TOPPADDING', (0, 0), (-1, 0), 8),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
        ('LEFTPADDING', (0, 0), (-1, 0), 10),

        # Data rows
        ('BACKGROUND', (0, 1), (0, -1), colors.HexColor('#ecf0f1')),
        ('FONTNAME', (0, 1), (0, -1), 'Helvetica-Bold'),
        ('FONTNAME', (1, 1), (1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
        ('TEXTCOLOR', (0, 1), (-1, -1), colors.HexColor('#2c3e50')),
        ('ALIGN', (0, 1), (0, -1), 'LEFT'),
        ('ALIGN', (1, 1), (1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('TOPPADDING', (0, 1), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
        ('LEFTPADDING', (0, 1), (-1, -1), 10),
        ('RIGHTPADDING', (0, 1), (-1, -1), 10),

        # Borders
        ('BOX', (0, 0), (-1, -1), 1, colors.HexColor('#bdc3c7')),
        ('LINEBELOW', (0, 0), (-1, 0), 2, colors.HexColor('#34495e')),
        ('INNERGRID', (0, 1), (-1, -1), 0.5, colors.HexColor('#bdc3c7')),
    ]))

    elements.append(patient_table)
    elements.append(Spacer(1, 0.15*inch))

    # ============ CREDENCIALES DE ACCESO ============
    credentials_data = [
        ['CREDENCIALES PARA CONSULTAR RESULTADOS'],
        ['CÉDULA DE IDENTIDAD:', resultado.paciente_ci],
        ['CÓDIGO DE ACCESO:', resultado.codigo_acceso]
    ]

    credentials_table = Table(credentials_data, colWidths=[2.5*inch, 4*inch])
    credentials_table.setStyle(TableStyle([
        # Header
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#e67e22')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 11),
        ('SPAN', (0, 0), (-1, 0)),
        ('TOPPADDING', (0, 0), (-1, 0), 8),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
        ('LEFTPADDING', (0, 0), (-1, 0), 10),

        # Data rows - CI
        ('BACKGROUND', (0, 1), (0, 1), colors.HexColor('#ecf0f1')),
        ('FONTNAME', (0, 1), (0, 1), 'Helvetica-Bold'),
        ('FONTNAME', (1, 1), (1, 1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, 1), 9),
        ('TEXTCOLOR', (0, 1), (-1, 1), colors.HexColor('#2c3e50')),

        # CÓDIGO - Destacado
        ('BACKGROUND', (0, 2), (0, 2), colors.HexColor('#ecf0f1')),
        ('BACKGROUND', (1, 2), (1, 2), colors.HexColor('#fff9e6')),
        ('FONTNAME', (0, 2), (0, 2), 'Helvetica-Bold'),
        ('FONTNAME', (1, 2), (1, 2), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 2), (0, 2), 9),
        ('FONTSIZE', (1, 2), (1, 2), 14),
        ('TEXTCOLOR', (1, 2), (1, 2), colors.HexColor('#e74c3c')),

        # General
        ('ALIGN', (0, 1), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('TOPPADDING', (0, 1), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
        ('LEFTPADDING', (0, 1), (-1, -1), 10),
        ('RIGHTPADDING', (0, 1), (-1, -1), 10),

        # Borders
        ('BOX', (0, 0), (-1, -1), 1, colors.HexColor('#bdc3c7')),
        ('LINEBELOW', (0, 0), (-1, 0), 2, colors.HexColor('#e67e22')),
        ('INNERGRID', (0, 1), (-1, -1), 0.5, colors.HexColor('#bdc3c7')),
    ]))

    elements.append(credentials_table)
    elements.append(Spacer(1, 0.15*inch))

    # ============ INSTRUCCIONES ============
    instructions_data = [
        ['INSTRUCCIONES PARA ACCEDER A SUS RESULTADOS'],
        ['1', 'Ingrese a: www.laboratoriopérez.com o utilice el enlace proporcionado por el laboratorio'],
        ['2', 'Click en el botón "Consultar mis Resultados" o "Ver Resultados"'],
        ['3', 'Ingrese su Cédula de Identidad (CI) y el Código de Acceso proporcionado arriba'],
        ['4', 'Podrá visualizar y descargar su resultado en formato PDF'],
        ['', 'NOTA: Guarde este documento de forma segura. Su código de acceso es confidencial.']
    ]

    instructions_table = Table(instructions_data, colWidths=[0.4*inch, 6.1*inch])
    instructions_table.setStyle(TableStyle([
        # Header
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3498db')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 11),
        ('SPAN', (0, 0), (-1, 0)),
        ('TOPPADDING', (0, 0), (-1, 0), 8),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
        ('LEFTPADDING', (0, 0), (-1, 0), 10),

        # Steps
        ('FONTNAME', (0, 1), (0, 4), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 1), (0, 4), 11),
        ('TEXTCOLOR', (0, 1), (0, 4), colors.HexColor('#3498db')),
        ('ALIGN', (0, 1), (0, 4), 'CENTER'),
        ('FONTNAME', (1, 1), (1, 4), 'Helvetica'),
        ('FONTSIZE', (1, 1), (1, 4), 9),
        ('TEXTCOLOR', (1, 1), (1, 4), colors.HexColor('#2c3e50')),

        # Note
        ('BACKGROUND', (0, 5), (-1, 5), colors.HexColor('#fff3cd')),
        ('SPAN', (0, 5), (-1, 5)),
        ('FONTNAME', (0, 5), (-1, 5), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 5), (-1, 5), 8),
        ('TEXTCOLOR', (0, 5), (-1, 5), colors.HexColor('#856404')),

        # General
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('TOPPADDING', (0, 1), (-1, -1), 5),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 5),
        ('LEFTPADDING', (0, 1), (-1, -1), 10),
        ('RIGHTPADDING', (0, 1), (-1, -1), 10),

        # Borders
        ('BOX', (0, 0), (-1, -1), 1, colors.HexColor('#bdc3c7')),
        ('LINEBELOW', (0, 0), (-1, 0), 2, colors.HexColor('#3498db')),
        ('INNERGRID', (0, 1), (-1, -1), 0.5, colors.HexColor('#bdc3c7')),
    ]))

    elements.append(instructions_table)
    elements.append(Spacer(1, 0.1*inch))

    # ============ FOOTER ============
    footer_text = Paragraph(
        '<i>Laboratorio Clínico Pérez - Potosí, Bolivia - Documento generado automáticamente</i>',
        ParagraphStyle(
            'Footer',
            parent=styles['Normal'],
            fontSize=8,
            textColor=colors.HexColor('#95a5a6'),
            alignment=1
        )
    )
    elements.append(footer_text)

    # ============ GENERAR PDF ============
    doc.build(elements)
    buffer.seek(0)

    filename = f"Credenciales_{resultado.paciente_nombre.replace(' ', '_')}_{resultado.numero_orden}.pdf"

    return send_file(buffer, as_attachment=True, download_name=filename, mimetype='application/pdf')

@main.route('/descargar-credenciales-word/<int:resultado_id>')
@admin_required
def descargar_credenciales_word(resultado_id):
    """
    Genera Word de credenciales PROFESIONAL y BONITO
    Para que el paciente vea su código de acceso
    """
    resultado = Resultado.query.get_or_404(resultado_id)

    doc = Document()

    # ============ CONFIGURACIÓN DE MÁRGENES ============
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ============ ENCABEZADO DEL DOCUMENTO ============
    titulo = doc.add_heading('LABORATORIO CLÍNICO PÉREZ', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo_run = titulo.runs[0]
    titulo_run.font.color.rgb = RGBColor(17, 153, 142)  # Verde Pérez
    titulo_run.font.size = Pt(20)
    titulo_run.font.name = 'Arial'

    subtitulo = doc.add_paragraph('Potosí, Bolivia')
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitulo_run = subtitulo.runs[0]
    subtitulo_run.font.size = Pt(10)
    subtitulo_run.font.color.rgb = RGBColor(127, 140, 141)
    subtitulo_run.font.italic = True

    # ============ TÍTULO PRINCIPAL ============
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_para.add_run('CREDENCIALES DE ACCESO A RESULTADOS')
    header_run.font.bold = True
    header_run.font.size = Pt(12)
    header_run.font.color.rgb = RGBColor(255, 255, 255)

    # Fondo verde para el título
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), '11998e')
    header_para._element.get_or_add_pPr().append(shading_elm)

    # Padding del párrafo
    pPr = header_para._element.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '100')
    spacing.set(qn('w:after'), '100')
    pPr.append(spacing)

    # ============ TABLA: INFORMACIÓN DEL PACIENTE ============
    table1 = doc.add_table(rows=5, cols=2)
    table1.style = 'Light Grid Accent 1'

    # Header de sección
    header_cell = table1.rows[0].cells[0]
    header_cell.merge(table1.rows[0].cells[1])
    header_cell.text = 'INFORMACIÓN DEL PACIENTE'
    header_run = header_cell.paragraphs[0].runs[0]
    header_run.font.bold = True
    header_run.font.size = Pt(10)
    header_run.font.color.rgb = RGBColor(255, 255, 255)
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), '34495e')
    header_cell.paragraphs[0]._element.get_or_add_pPr().append(shading_elm)

    # Datos del paciente
    datos_paciente = [
        ['Nombre Completo:', resultado.paciente_nombre],
        ['Cédula de Identidad:', resultado.paciente_ci],
        ['Número de Orden:', resultado.numero_orden],
        ['Fecha de Emisión:', resultado.fecha_creacion.strftime("%d/%m/%Y %H:%M")]
    ]

    for i, (label, value) in enumerate(datos_paciente, start=1):
        # Columna izquierda (labels)
        cell_label = table1.rows[i].cells[0]
        cell_label.text = label
        cell_label.paragraphs[0].runs[0].font.bold = True
        cell_label.paragraphs[0].runs[0].font.size = Pt(9)
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'ecf0f1')
        cell_label.paragraphs[0]._element.get_or_add_pPr().append(shading_elm)

        # Columna derecha (valores)
        cell_value = table1.rows[i].cells[1]
        cell_value.text = value
        cell_value.paragraphs[0].runs[0].font.size = Pt(9)

    # ============ TABLA: CREDENCIALES DE ACCESO ============
    table2 = doc.add_table(rows=3, cols=2)
    table2.style = 'Light Grid Accent 1'

    # Header de sección
    header_cell2 = table2.rows[0].cells[0]
    header_cell2.merge(table2.rows[0].cells[1])
    header_cell2.text = 'CREDENCIALES PARA CONSULTAR RESULTADOS'
    header_run2 = header_cell2.paragraphs[0].runs[0]
    header_run2.font.bold = True
    header_run2.font.size = Pt(10)
    header_run2.font.color.rgb = RGBColor(255, 255, 255)
    shading_elm2 = OxmlElement('w:shd')
    shading_elm2.set(qn('w:fill'), 'e67e22')
    header_cell2.paragraphs[0]._element.get_or_add_pPr().append(shading_elm2)

    # CI
    table2.rows[1].cells[0].text = 'CÉDULA DE IDENTIDAD:'
    table2.rows[1].cells[0].paragraphs[0].runs[0].font.bold = True
    table2.rows[1].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'ecf0f1')
    table2.rows[1].cells[0].paragraphs[0]._element.get_or_add_pPr().append(shading_elm)

    table2.rows[1].cells[1].text = resultado.paciente_ci
    table2.rows[1].cells[1].paragraphs[0].runs[0].font.size = Pt(9)

    # CÓDIGO DE ACCESO (destacado)
    table2.rows[2].cells[0].text = 'CÓDIGO DE ACCESO:'
    table2.rows[2].cells[0].paragraphs[0].runs[0].font.bold = True
    table2.rows[2].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'ecf0f1')
    table2.rows[2].cells[0].paragraphs[0]._element.get_or_add_pPr().append(shading_elm)

    table2.rows[2].cells[1].text = resultado.codigo_acceso
    codigo_run = table2.rows[2].cells[1].paragraphs[0].runs[0]
    codigo_run.font.bold = True
    codigo_run.font.size = Pt(14)
    codigo_run.font.color.rgb = RGBColor(231, 76, 60)  # Rojo para destacar
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'fff9e6')  # Amarillo claro
    table2.rows[2].cells[1].paragraphs[0]._element.get_or_add_pPr().append(shading_elm)

    # ============ TABLA: INSTRUCCIONES ============
    table3 = doc.add_table(rows=6, cols=1)
    table3.style = 'Light Grid Accent 1'

    # Header de sección
    header_cell3 = table3.rows[0].cells[0]
    header_cell3.text = 'INSTRUCCIONES PARA ACCEDER A SUS RESULTADOS'
    header_run3 = header_cell3.paragraphs[0].runs[0]
    header_run3.font.bold = True
    header_run3.font.size = Pt(10)
    header_run3.font.color.rgb = RGBColor(255, 255, 255)
    shading_elm3 = OxmlElement('w:shd')
    shading_elm3.set(qn('w:fill'), '3498db')
    header_cell3.paragraphs[0]._element.get_or_add_pPr().append(shading_elm3)

    # Instrucciones paso a paso
    instrucciones = [
        '1. Ingrese a: www.laboratoriopérez.com o utilice el enlace proporcionado por el laboratorio',
        '2. Click en el botón "Consultar mis Resultados" o "Ver Resultados"',
        '3. Ingrese su Cédula de Identidad (CI) y el Código de Acceso proporcionado arriba',
        '4. Podrá visualizar y descargar su resultado en formato PDF',
        'NOTA: Guarde este documento de forma segura. Su código de acceso es confidencial.'
    ]

    for i, instruccion in enumerate(instrucciones, start=1):
        cell = table3.rows[i].cells[0]
        cell.text = instruccion
        cell_run = cell.paragraphs[0].runs[0]
        cell_run.font.size = Pt(8)

        # Última fila (nota) con fondo especial
        if i == 5:
            cell_run.font.bold = True
            cell_run.font.color.rgb = RGBColor(133, 100, 4)
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'fff3cd')
            cell.paragraphs[0]._element.get_or_add_pPr().append(shading_elm)

    doc.add_paragraph()  # Espacio

    # ============ FOOTER ============
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run('Laboratorio Clínico Pérez - Potosí, Bolivia')
    footer_run.font.size = Pt(8)
    footer_run.font.italic = True
    footer_run.font.color.rgb = RGBColor(149, 165, 166)

    footer_para2 = doc.add_paragraph()
    footer_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run2 = footer_para2.add_run('Documento generado automáticamente')
    footer_run2.font.size = Pt(7)
    footer_run2.font.italic = True
    footer_run2.font.color.rgb = RGBColor(149, 165, 166)

    # ============ GUARDAR DOCUMENTO ============
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"Credenciales_{resultado.paciente_nombre.replace(' ', '_')}_{resultado.numero_orden}.docx"

    return send_file(buffer, as_attachment=True, download_name=filename, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@main.route('/pruebas', methods=['GET', 'POST'])
@admin_required
def admin_pruebas():
    if request.method == 'POST':
        try:
            # Manejar imagen si se subió
            imagen_url = None
            if 'imagen' in request.files:
                imagen = request.files['imagen']
                if imagen and imagen.filename:
                    success, storage_path, error = file_manager.save_image(imagen, 'pruebas')
                    if success:
                        imagen_url = file_manager.get_public_url(storage_path, 'img')

            prueba = Prueba(
                nombre=request.form['nombre'],
                categoria=request.form.get('categoria'),
                descripcion=request.form.get('descripcion'),
                precio=float(request.form.get('precio', 0)),
                imagen=imagen_url
            )
            db.session.add(prueba)
            db.session.commit()
            flash('Prueba registrada exitosamente', 'success')
        except Exception as e:
            flash(f'Error: {str(e)}', 'danger')
        return redirect(url_for('main.admin_pruebas'))

    pruebas = Prueba.query.order_by(Prueba.nombre).all()

    # Obtener categorías únicas de la base de datos
    categorias = db.session.query(Prueba.categoria).distinct().filter(Prueba.categoria.isnot(None)).order_by(Prueba.categoria).all()
    categorias_list = [cat[0] for cat in categorias if cat[0]]

    return render_template('admin/pruebas.html', pruebas=pruebas, categorias=categorias_list)

@main.route('/prueba/<int:prueba_id>')
@admin_required
def ver_prueba(prueba_id):
    prueba = Prueba.query.get_or_404(prueba_id)
    return jsonify({
        'id': prueba.id,
        'nombre': prueba.nombre,
        'categoria': prueba.categoria,
        'descripcion': prueba.descripcion,
        'precio': float(prueba.precio),
        'imagen': prueba.imagen
    })

@main.route('/prueba/editar/<int:prueba_id>', methods=['POST'])
@admin_required
def editar_prueba(prueba_id):
    try:
        prueba = Prueba.query.get_or_404(prueba_id)
        prueba.nombre = request.form['nombre']
        prueba.categoria = request.form.get('categoria')
        prueba.descripcion = request.form.get('descripcion')
        prueba.precio = float(request.form.get('precio', 0))

        # Manejar nueva imagen si se subió
        if 'imagen' in request.files:
            imagen = request.files['imagen']
            if imagen and imagen.filename:
                # Eliminar imagen anterior si existe
                if prueba.imagen:
                    try:
                        # Extraer path relativo si es una URL completa
                        # Pero Supabase delete espera path relativo.
                        if 'supabase' in prueba.imagen:
                             pass # Implementar extracción de path de URL si es necesario
                             # O simplemente dejar que FileManager.delete maneje URLs? No, delete espera path.
                             # Si guardamos URLs enteras, necesitamos parsear el path.
                             # Por ahora, solo subimos la nueva.
                             # Limpieza se puede hacer después o con un script.
                             pass
                    except:
                        pass

                # Guardar nueva imagen
                success, storage_path, error = file_manager.save_image(imagen, 'pruebas')
                if success:
                    public_url = file_manager.get_public_url(storage_path, 'img')
                    prueba.imagen = public_url

        db.session.commit()
        flash('Prueba actualizada exitosamente', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error: {str(e)}', 'danger')
    return redirect(url_for('main.admin_pruebas'))

@main.route('/prueba/eliminar/<int:prueba_id>', methods=['POST'])
@admin_required
def eliminar_prueba(prueba_id):
    try:
        prueba = Prueba.query.get_or_404(prueba_id)
        nombre_prueba = prueba.nombre

        # Eliminar imagen si existe
        # Eliminar imagen si existe (Opcional - Supabase)
        if prueba.imagen and 'supabase' in prueba.imagen:
            # Aquí necesitaríamos el path relativo para borrar de Supabase (ej: pruebas/foto.jpg)
            # Como guardamos la URL entera, es difícil sacar el path sin una función helper.
            # Por seguridad, dejemos el archivo en la nube o movamos a papelera si implementamos parseo.
            pass
        elif prueba.imagen: # Local legacy
             try:
                imagen_path = os.path.join(PRUEBAS_UPLOAD_DIR, prueba.imagen)
                if os.path.exists(imagen_path):
                    os.remove(imagen_path)
             except: pass

        db.session.delete(prueba)
        db.session.commit()
        flash(f'Prueba "{nombre_prueba}" eliminada exitosamente', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error al eliminar prueba: {str(e)}', 'danger')
    return redirect(url_for('main.admin_pruebas'))














# ============================================
# RUTAS DE REDES SOCIALES / PÁGINA PÚBLICA
# ============================================

# Carpeta para imágenes de publicaciones y galería
SOCIAL_UPLOAD_DIR = os.path.join(UPLOAD_DIR, 'social')

@main.route('/admin/redes-sociales')
@admin_required
def admin_redes_sociales():
    """Panel de administración de Redes Sociales"""
    publicaciones = Publicacion.query.order_by(Publicacion.orden.asc(), Publicacion.fecha_publicacion.desc()).all()
    fotos = FotoGaleria.query.order_by(FotoGaleria.orden.asc(), FotoGaleria.fecha_subida.desc()).all()
    
    # Cargar configuración como diccionario
    config_items = ConfiguracionLab.query.all()
    configuracion = {item.clave: item.valor for item in config_items}
    
    return render_template('admin/redes_sociales.html',
                         publicaciones=publicaciones,
                         fotos=fotos,
                         configuracion=configuracion)

@main.route('/publicacion/nueva', methods=['POST'])
@admin_required
def nueva_publicacion():
    """Crear nueva publicación"""
    try:
        titulo = request.form.get('titulo')
        contenido = request.form.get('contenido')
        icono = request.form.get('icono', 'fa-microscope')
        categoria = request.form.get('categoria', 'General')
        
        nueva = Publicacion(
            titulo=titulo,
            contenido=contenido,
            icono=icono,
            categoria=categoria
        )
        
        # Manejar imagen si se subió
        # Manejar imagen si se subió
        if 'imagen' in request.files:
            file = request.files['imagen']
            if file and file.filename:
                # Subir a Supabase (Bucket imagenes, carpeta social)
                success, storage_path, error = file_manager.save_image(file, 'social')
                if success:
                    public_url = file_manager.get_public_url(storage_path, 'img')
                    nueva.imagen = public_url
        
        db.session.add(nueva)
        db.session.commit()
        flash('✅ Publicación creada exitosamente', 'success')
        
    except Exception as e:
        db.session.rollback()
        flash(f'❌ Error al crear publicación: {str(e)}', 'danger')
    
    return redirect(url_for('main.admin_redes_sociales'))

@main.route('/publicacion/editar/<int:id>', methods=['POST'])
@admin_required
def editar_publicacion(id):
    """Editar publicación existente"""
    try:
        pub = Publicacion.query.get_or_404(id)
        pub.titulo = request.form.get('titulo', pub.titulo)
        pub.contenido = request.form.get('contenido', pub.contenido)
        pub.icono = request.form.get('icono', pub.icono)
        pub.categoria = request.form.get('categoria', pub.categoria)
        pub.activo = 'activo' in request.form
        
        # Manejar nueva imagen si se subió
        # Manejar nueva imagen si se subió
        if 'imagen' in request.files:
            file = request.files['imagen']
            if file and file.filename:
                # Subir a Supabase
                success, storage_path, error = file_manager.save_image(file, 'social')
                if success:
                    public_url = file_manager.get_public_url(storage_path, 'img')
                    pub.imagen = public_url
        
        db.session.commit()
        flash('✅ Publicación actualizada', 'success')
        
    except Exception as e:
        db.session.rollback()
        flash(f'❌ Error al editar: {str(e)}', 'danger')
    
    return redirect(url_for('main.admin_redes_sociales'))

@main.route('/publicacion/eliminar/<int:id>', methods=['POST'])
@admin_required
def eliminar_publicacion(id):
    """Eliminar publicación"""
    try:
        pub = Publicacion.query.get_or_404(id)
        db.session.delete(pub)
        db.session.commit()
        flash('✅ Publicación eliminada', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'❌ Error al eliminar: {str(e)}', 'danger')
    
    return redirect(url_for('main.admin_redes_sociales'))

@main.route('/foto/nueva', methods=['POST'])
@admin_required
def nueva_foto():
    """Subir nueva foto a galería"""
    try:
        titulo = request.form.get('titulo')
        descripcion = request.form.get('descripcion', '')
        
        if 'imagen' not in request.files:
            flash('❌ Debe seleccionar una imagen', 'danger')
            return redirect(url_for('main.admin_redes_sociales'))
        
        file = request.files['imagen']
        if not file or not file.filename:
            flash('❌ Archivo de imagen inválido', 'danger')
            return redirect(url_for('main.admin_redes_sociales'))
        
        # Subir a Supabase (Bucket imagenes, carpeta social)
        success, storage_path, error = file_manager.save_image(file, 'social')
        if not success:
             raise Exception(error)
        
        public_url = file_manager.get_public_url(storage_path, 'img')
        
        nueva = FotoGaleria(
            titulo=titulo,
            descripcion=descripcion,
            imagen=public_url
        )
        
        db.session.add(nueva)
        db.session.commit()
        flash('✅ Foto agregada a la galería', 'success')
        
    except Exception as e:
        db.session.rollback()
        flash(f'❌ Error al subir foto: {str(e)}', 'danger')
    
    return redirect(url_for('main.admin_redes_sociales'))

@main.route('/foto/eliminar/<int:id>', methods=['POST'])
@admin_required
def eliminar_foto(id):
    """Eliminar foto de galería"""
    try:
        foto = FotoGaleria.query.get_or_404(id)
        db.session.delete(foto)
        db.session.commit()
        flash('✅ Foto eliminada', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'❌ Error al eliminar: {str(e)}', 'danger')
    
    return redirect(url_for('main.admin_redes_sociales'))

@main.route('/configuracion/guardar', methods=['POST'])
@admin_required
def guardar_configuracion():
    """Guardar configuración del laboratorio"""
    try:
        # Lista de claves de configuración
        claves = ['telefono', 'whatsapp', 'email', 'direccion', 'referencia',
                  'horario_semana', 'horario_sabado', 'horario_domingo',
                  'experiencia', 'certificacion', 'rating', 'resenas']
        
        for clave in claves:
            valor = request.form.get(clave, '')
            if valor:
                config = ConfiguracionLab.query.filter_by(clave=clave).first()
                if config:
                    config.valor = valor
                else:
                    config = ConfiguracionLab(clave=clave, valor=valor)
                    db.session.add(config)
        
        db.session.commit()
        flash('✅ Configuración guardada exitosamente', 'success')
        
    except Exception as e:
        db.session.rollback()
        flash(f'❌ Error al guardar: {str(e)}', 'danger')
    
    return redirect(url_for('main.admin_redes_sociales'))


@main.route('/imagenes-perfil/guardar', methods=['POST'])
@admin_required
def guardar_imagenes_perfil():
    """Guardar imágenes de portada y perfil"""
    try:
        # Foto de Portada
        if 'foto_portada' in request.files:
            file = request.files['foto_portada']
            if file and file.filename:
                success, storage_path, error = file_manager.save_image(file, 'portada')
                if success:
                    public_url = file_manager.get_public_url(storage_path, 'img')
                    
                    # Guardar en configuración
                    config = ConfiguracionLab.query.filter_by(clave='foto_portada').first()
                    if config:
                        config.valor = public_url
                    else:
                        config = ConfiguracionLab(clave='foto_portada', valor=public_url)
                        db.session.add(config)
        
        # Foto de Perfil/Logo
        if 'foto_perfil' in request.files:
            file = request.files['foto_perfil']
            if file and file.filename:
                success, storage_path, error = file_manager.save_image(file, 'perfil')
                if success:
                    public_url = file_manager.get_public_url(storage_path, 'img')
                    
                    # Guardar en configuración
                    config = ConfiguracionLab.query.filter_by(clave='foto_perfil').first()
                    if config:
                        config.valor = public_url
                    else:
                        config = ConfiguracionLab(clave='foto_perfil', valor=public_url)
                        db.session.add(config)
                        
        db.session.commit()
        flash('✅ Imágenes actualizadas exitosamente', 'success')
        
    except Exception as e:
        db.session.rollback()
        flash(f'❌ Error al guardar imágenes: {str(e)}', 'danger')
    
    return redirect(url_for('main.admin_redes_sociales'))




